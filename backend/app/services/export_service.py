"""
Export Service: Markdown → DOCX / PDF conversion for literature reviews.

Uses python-docx to create a well-formatted Word document from
the review's Markdown content, preserving:
- Heading hierarchy (H1, H2, H3)
- Bold/italic inline formatting
- Numbered reference lists
- Basic paragraph structure

Also supports Markdown → HTML → PDF conversion using xhtml2pdf.
"""

import io
import re
import logging
from typing import Optional, Tuple, List

import markdown as md_lib

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

logger = logging.getLogger(__name__)


class MarkdownToDocxConverter:
    """Convert a Markdown string into a python-docx Document."""

    def __init__(self):
        self._doc: Optional[Document] = None

    def convert(self, markdown_text: str, title: str = "Literature Review") -> Document:
        """
        Main entry point: parse Markdown and produce a Document object.

        Args:
            markdown_text: The full Markdown content of the review.
            title: Fallback title if not found in Markdown.

        Returns:
            A python-docx Document ready to be saved.
        """
        self._doc = Document()
        self._setup_styles()

        lines = markdown_text.split("\n")
        i = 0
        in_references = False

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip empty lines
            if not stripped:
                i += 1
                continue

            # Horizontal rule
            if stripped in ("---", "***", "___"):
                # self._doc.add_page_break()
                i += 1
                continue

            # Headings
            if stripped.startswith("# "):
                heading_text = stripped[2:].strip()
                self._add_heading(heading_text, level=1)
                i += 1
                continue

            if stripped.startswith("## "):
                heading_text = stripped[3:].strip()
                # Check if entering References section
                if heading_text.lower() in ("references", "参考文献"):
                    in_references = True
                self._add_heading(heading_text, level=2)
                i += 1
                continue

            if stripped.startswith("### "):
                heading_text = stripped[4:].strip()
                self._add_heading(heading_text, level=3)
                i += 1
                continue

            # Blockquote (used for abstract description)
            if stripped.startswith("> "):
                quote_text = stripped[2:].strip()
                p = self._doc.add_paragraph(style="Quote")
                self._add_formatted_text(p, quote_text)
                i += 1
                continue

            # Reference list items (numbered: "1. Author..." or "[1] Author...")
            if in_references and (re.match(r'^\d+\.\s', stripped) or re.match(r'^\[\d+\]\s', stripped)):
                p = self._doc.add_paragraph(style="List Number")
                # Remove the number prefix
                ref_text = re.sub(r'^(\d+\.\s|\[\d+\]\s)', '', stripped)
                self._add_formatted_text(p, ref_text)
                i += 1
                continue

            # Bullet list items
            if stripped.startswith("- ") or stripped.startswith("* "):
                bullet_text = stripped[2:].strip()
                p = self._doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(p, bullet_text)
                i += 1
                continue

            # Regular paragraph — collect consecutive non-empty, non-special lines
            para_lines = [stripped]
            i += 1
            while i < len(lines):
                next_line = lines[i].strip()
                if (not next_line or
                    next_line.startswith("#") or
                    next_line.startswith("> ") or
                    next_line.startswith("- ") or
                    next_line.startswith("* ") or
                    next_line in ("---", "***", "___") or
                    (in_references and re.match(r'^(\d+\.\s|\[\d+\]\s)', next_line))):
                    break
                para_lines.append(next_line)
                i += 1

            full_para = " ".join(para_lines)
            p = self._doc.add_paragraph()
            self._add_formatted_text(p, full_para)

        return self._doc

    def convert_to_bytes(self, markdown_text: str, title: str = "Literature Review") -> bytes:
        """Convert Markdown to DOCX and return as bytes (for HTTP response)."""
        doc = self.convert(markdown_text, title)
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    def _setup_styles(self):
        """Configure document styles for academic formatting."""
        style = self._doc.styles

        # Normal style
        normal = style["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.5

        # Heading 1
        h1 = style["Heading 1"]
        h1.font.name = "Times New Roman"
        h1.font.size = Pt(16)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(0, 0, 0)
        h1.paragraph_format.space_before = Pt(12)
        h1.paragraph_format.space_after = Pt(8)
        h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Heading 2
        h2 = style["Heading 2"]
        h2.font.name = "Times New Roman"
        h2.font.size = Pt(14)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(0, 0, 0)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(6)

        # Heading 3
        h3 = style["Heading 3"]
        h3.font.name = "Times New Roman"
        h3.font.size = Pt(13)
        h3.font.bold = True
        h3.font.italic = True
        h3.font.color.rgb = RGBColor(0, 0, 0)
        h3.paragraph_format.space_before = Pt(8)
        h3.paragraph_format.space_after = Pt(4)

        # Set page margins
        for section in self._doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def _add_heading(self, text: str, level: int = 1):
        """Add a heading with the specified level."""
        heading = self._doc.add_heading(level=level)
        # Remove any markdown formatting from heading text
        clean_text = re.sub(r'[*_`]', '', text)
        heading.text = clean_text

    def _add_formatted_text(self, paragraph, text: str):
        """
        Parse inline Markdown formatting and add runs to a paragraph.
        Supports: **bold**, *italic*, ***bold italic***, `code`
        """
        # Pattern to match inline formatting
        # Order matters: bold+italic first, then bold, then italic, then code
        pattern = re.compile(
            r'(\*\*\*(.+?)\*\*\*)'   # ***bold italic***
            r'|(\*\*(.+?)\*\*)'       # **bold**
            r'|(\*(.+?)\*)'           # *italic*
            r'|(`(.+?)`)'             # `code`
        )

        last_end = 0
        for match in pattern.finditer(text):
            # Add text before this match
            if match.start() > last_end:
                paragraph.add_run(text[last_end:match.start()])

            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
            elif match.group(4):  # **bold**
                run = paragraph.add_run(match.group(4))
                run.bold = True
            elif match.group(6):  # *italic*
                run = paragraph.add_run(match.group(6))
                run.italic = True
            elif match.group(8):  # `code`
                run = paragraph.add_run(match.group(8))
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(80, 80, 80)

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            paragraph.add_run(text[last_end:])


def export_review_to_docx(content: str, title: str = "Literature Review") -> bytes:
    """
    Convenience function: convert review Markdown to DOCX bytes.

    Args:
        content: Full Markdown content of the review.
        title: Review title (used as fallback).

    Returns:
        DOCX file content as bytes.
    """
    converter = MarkdownToDocxConverter()
    return converter.convert_to_bytes(content, title)


# ============================================================
# PDF Export: Markdown → HTML → PDF (via xhtml2pdf)
# ============================================================

# Academic PDF stylesheet
_PDF_CSS = """
@page {
    size: A4;
    margin: 2.5cm 2.5cm 2.5cm 2.5cm;
    @bottom-center {
        content: "Page " counter(page) " of " counter(pages);
        font-size: 9pt;
        color: #666;
        font-family: "Times New Roman", Times, serif;
    }
}

body {
    font-family: "Times New Roman", Times, "SimSun", "宋体", serif;
    font-size: 12pt;
    line-height: 1.6;
    color: #000;
}

h1 {
    font-size: 18pt;
    font-weight: bold;
    text-align: center;
    margin-top: 12pt;
    margin-bottom: 10pt;
    page-break-after: avoid;
}

h2 {
    font-size: 15pt;
    font-weight: bold;
    margin-top: 14pt;
    margin-bottom: 8pt;
    border-bottom: 1px solid #ccc;
    padding-bottom: 3pt;
    page-break-after: avoid;
}

h3 {
    font-size: 13pt;
    font-weight: bold;
    font-style: italic;
    margin-top: 10pt;
    margin-bottom: 6pt;
    page-break-after: avoid;
}

p {
    text-align: justify;
    margin-bottom: 6pt;
    text-indent: 0;
}

blockquote {
    margin-left: 1cm;
    margin-right: 1cm;
    padding-left: 0.5cm;
    border-left: 3px solid #ccc;
    font-style: italic;
    color: #333;
}

ul, ol {
    margin-left: 1cm;
    margin-bottom: 6pt;
}

li {
    margin-bottom: 3pt;
}

code {
    font-family: "Consolas", "Courier New", monospace;
    font-size: 10pt;
    background-color: #f5f5f5;
    padding: 1px 3px;
}

hr {
    border: none;
    border-top: 1px solid #ccc;
    margin: 12pt 0;
}

/* Reference list styling */
.references p, .references li {
    text-indent: -1cm;
    padding-left: 1cm;
    font-size: 11pt;
}
"""


class MarkdownToPdfConverter:
    """Convert a Markdown string into a PDF via HTML intermediate format."""

    def convert(self, markdown_text: str, title: str = "Literature Review") -> bytes:
        """
        Convert Markdown content to PDF bytes.

        Args:
            markdown_text: The full Markdown content of the review.
            title: Document title for the HTML head.

        Returns:
            PDF file content as bytes.
        """
        # Step 1: Convert Markdown to HTML
        html_body = md_lib.markdown(
            markdown_text,
            extensions=["extra", "smarty", "sane_lists"],
            output_format="html5",
        )

        # Step 2: Wrap references section with a class for styling
        html_body = self._tag_references_section(html_body)

        # Step 3: Build full HTML document
        html_doc = self._build_html_document(html_body, title)

        # Step 4: Convert HTML to PDF
        return self._html_to_pdf(html_doc)

    def _tag_references_section(self, html: str) -> str:
        """
        Wrap the References/参考文献 section in a <div class="references">
        for targeted CSS styling (hanging indent etc.).
        """
        # Find the references heading and wrap everything after it
        patterns = [
            (r'(<h2>References</h2>)', r'<div class="references">\1'),
            (r'(<h2>参考文献</h2>)', r'<div class="references">\1'),
        ]
        for pattern, replacement in patterns:
            if re.search(pattern, html, re.IGNORECASE):
                html = re.sub(pattern, replacement, html, flags=re.IGNORECASE)
                html += "</div>"
                break
        return html

    def _build_html_document(self, body: str, title: str) -> str:
        """Build a complete HTML5 document with embedded CSS."""
        return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="utf-8"/>
    <title>{title}</title>
    <style>
    {_PDF_CSS}
    </style>
</head>
<body>
{body}
</body>
</html>"""

    def _html_to_pdf(self, html: str) -> bytes:
        """
        Convert HTML string to PDF bytes using xhtml2pdf.
        
        Falls back to a simple error message if xhtml2pdf is not installed.
        """
        try:
            from xhtml2pdf import pisa
        except ImportError:
            raise RuntimeError(
                "xhtml2pdf is not installed. "
                "Install it with: pip install xhtml2pdf"
            )

        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(
            src=html,
            dest=buffer,
            encoding="utf-8",
        )

        if pisa_status.err:
            logger.error(f"xhtml2pdf reported {pisa_status.err} error(s) during PDF generation")

        buffer.seek(0)
        return buffer.read()


def export_review_to_pdf(content: str, title: str = "Literature Review") -> bytes:
    """
    Convenience function: convert review Markdown to PDF bytes.

    Args:
        content: Full Markdown content of the review.
        title: Review title (used as fallback).

    Returns:
        PDF file content as bytes.
    """
    converter = MarkdownToPdfConverter()
    return converter.convert(content, title)